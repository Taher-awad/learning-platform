from fastapi import FastAPI, UploadFile, File, HTTPException
import boto3
from pydantic_settings import BaseSettings, SettingsConfigDict
import os
import uuid
import json
from confluent_kafka import Producer
import io
from pypdf import PdfReader
import docx
from sqlalchemy.orm import Session
from fastapi import Depends
from database import init_db, get_db, Document
import google.generativeai as genai

class Settings(BaseSettings):
    S3_BUCKET_NAME: str = "learning-platform-document-storage-dev"
    AWS_REGION: str = "us-east-1"
    KAFKA_BOOTSTRAP_SERVERS: str = "kafka:29092"
    GOOGLE_API_KEY: str
    POSTGRES_USER: str = "user"
    POSTGRES_PASSWORD: str = "password"
    POSTGRES_DB: str = "db"
    POSTGRES_HOST: str = "postgres"
    
    model_config = SettingsConfigDict(env_file=".env", extra="ignore")

settings = Settings()

# Configure Gemini
genai.configure(api_key=settings.GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash')

app = FastAPI()

# S3 Client
s3_client = boto3.client('s3', region_name=settings.AWS_REGION)

# Kafka Producer
producer_conf = {'bootstrap.servers': settings.KAFKA_BOOTSTRAP_SERVERS}
producer = Producer(producer_conf)

def delivery_report(err, msg):
    if err is not None:
        print(f'Message delivery failed: {err}')
    else:
        print(f'Message delivered to {msg.topic()} [{msg.partition()}]')

@app.on_event("startup")
def startup_event():
    init_db()
    # Ensure bucket exists (optional check, usually assumed in prod)
    try:
        s3_client.head_bucket(Bucket=settings.S3_BUCKET_NAME)
    except Exception as e:
        print(f"Warning: Bucket {settings.S3_BUCKET_NAME} check failed: {e}")

@app.post("/api/documents/upload")
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    try:
        file_id = str(uuid.uuid4())
        file_extension = file.filename.split(".")[-1]
        object_name = f"{file_id}.{file_extension}"
        
        # Read file content
        content = await file.read()
        file_size = len(content)
        
        # Create DB Record
        db_doc = Document(
            id=file_id,
            filename=file.filename,
            s3_bucket=settings.S3_BUCKET_NAME,
            s3_key=object_name,
            size=file_size,
            status="uploading"
        )
        db.add(db_doc)
        db.commit()
        
        # Upload to S3
        s3_client.upload_fileobj(
            io.BytesIO(content),
            settings.S3_BUCKET_NAME,
            object_name
        )
        
        # Produce document.uploaded
        producer.produce(
            "document.uploaded",
            key=file_id,
            value=json.dumps({"document_id": file_id, "filename": file.filename}),
            callback=delivery_report
        )
        
        # Extract text
        text_content = ""
        if file_extension.lower() == "pdf":
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
        elif file_extension.lower() == "docx":
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        elif file_extension.lower() == "txt":
            text_content = content.decode("utf-8")
            
        # Update DB status
        db_doc.status = "processed"
        db.commit()
            
        # Produce document.processed
        producer.produce(
            "document.processed",
            key=file_id,
            value=json.dumps({
                "document_id": file_id,
                "text_content": text_content
            }),
            callback=delivery_report
        )
        
        # Generate Notes
        try:
            prompt = f"Generate concise study notes for the following text:\n\n{text_content[:10000]}" # Limit context
            response = model.generate_content(prompt)
            notes = response.text
            
            db_doc.notes = notes
            db.commit()
            
            # Produce notes.generated
            producer.produce(
                "notes.generated",
                key=file_id,
                value=json.dumps({
                    "document_id": file_id,
                    "notes": notes
                }),
                callback=delivery_report
            )
        except Exception as e:
            print(f"Error generating notes: {e}")
            # Non-blocking error for notes
        
        producer.flush()
        
        return {"id": file_id, "message": "Document processed and notes generated"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/documents")
def list_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).all()
    return docs

@app.get("/api/documents/{file_id}")
def get_document(file_id: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == file_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@app.delete("/api/documents/{file_id}")
def delete_document(file_id: str, db: Session = Depends(get_db)):
    try:
        doc = db.query(Document).filter(Document.id == file_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
            
        # Delete from S3
        s3_client.delete_object(Bucket=settings.S3_BUCKET_NAME, Key=doc.s3_key)
        
        # Delete from DB
        db.delete(doc)
        db.commit()
        
        return {"message": f"Document {file_id} deleted"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
def read_root():
    return {"message": "Document Service is running"}

@app.get("/api/documents/")
def read_root_prefix():
    # This handles the specific gateway test case and health check
    return {"message": "Document Service is running"}
